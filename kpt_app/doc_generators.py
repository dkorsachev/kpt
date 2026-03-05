import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class DocumentGenerator:
    """Класс для генерации документов на основе шаблонов"""
    
    def __init__(self):
        self.markers = {
            'kadastr_number': '##KADASTR_NUMBER##',
            'contract_number': '##CONTRACT_NUMBER##',
            'contract_date': '##CONTRACT_DATE##',
            'contract_date_short': '##CONTRACT_DATE_SHORT##',
            'current_date': '##CURRENT_DATE##',
            'current_date_short': '##CURRENT_DATE_SHORT##',
            'table_place': '##TABLE_PLACE##',
        }
    
    def format_date_for_doc(self, date_obj):
        """Форматирование даты для документа (полный формат)"""
        if isinstance(date_obj, str):
            return date_obj
        months = [
            'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
            'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
        ]
        return f"{date_obj.day} {months[date_obj.month-1]} {date_obj.year} года"
    
    def format_date_short(self, date_obj):
        """Форматирование даты в короткий формат ДД.ММ.ГГГГ"""
        if isinstance(date_obj, str):
            return date_obj
        return date_obj.strftime('%d.%m.%Y')
    
    def set_run_font_size(self, run, size=12):
        """Установка размера шрифта для run"""
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    def normalize_paragraph_font_size(self, paragraph, target_size=12, exclude_title=True):
        """Приведение шрифта параграфа к целевому размеру без потери содержимого"""
        # Пропускаем заголовок если нужно
        if exclude_title and "СЛУЖЕБНАЯ ЗАПИСКА" in paragraph.text:
            return
        
        # Пропускаем пустые параграфы
        if not paragraph.text.strip() and not paragraph.runs:
            return
        
        # Для каждого run в параграфе устанавливаем размер шрифта
        for run in paragraph.runs:
            # Проверяем, не является ли run изображением
            # Изображения обычно имеют элемент drawing
            has_drawing = False
            for child in run._element:
                if child.tag.endswith('drawing'):
                    has_drawing = True
                    break
            
            if not has_drawing:
                # Это текстовый run - меняем размер шрифта
                self.set_run_font_size(run, target_size)
    
    def normalize_table_font_size(self, table, target_size=12):
        """Приведение шрифта таблицы к целевому размеру"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self.set_run_font_size(run, target_size)
    
    def replace_text_in_paragraph(self, paragraph, old_text, new_text):
        """Замена текста в параграфе с сохранением форматирования"""
        if old_text not in paragraph.text:
            return False
        
        # Сохраняем все runs
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True
        
        return False
    
    def replace_in_document(self, doc, replacements):
        """Замена всех маркеров в документе"""
        for key, value in replacements.items():
            # Замена в параграфах
            for paragraph in doc.paragraphs:
                self.replace_text_in_paragraph(paragraph, key, value)
            
            # Замена в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, key, value)
    
    def set_cell_font_size_12(self, cell, bold=False):
        """Установка шрифта 12 для ячейки таблицы"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = bold
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            
            paragraph.paragraph_format.line_spacing = Pt(14)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
    
    def create_table_at_marker_sluzhebka(self, doc, marker_paragraph, rows_data):
        """Создание таблицы для служебной записки (шрифт 12)"""
        parent = marker_paragraph._element.getparent()
        
        table = doc.add_table(rows=1 + len(rows_data), cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Выравнивание по центру
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is not None:
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            tblPr.append(jc)
        
        if len(table.columns) >= 4:
            table.columns[0].width = Cm(1.2)
            table.columns[1].width = Cm(3.5)
            table.columns[2].width = Cm(4.5)
            table.columns[3].width = Cm(5.5)
        
        # Заголовки таблицы (шрифт 12, жирный)
        headers = ['№ п/п', 'Тип документа', 'Кадастровый номер', 'Основание']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            self.set_cell_font_size_12(cell, bold=True)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Данные таблицы (шрифт 12)
        for i, row in enumerate(rows_data, 1):
            cad_num = row['cadastral_number']
            
            if re.match(r'^\d{2}:\d{2}:\d{7}$', cad_num):
                doc_type = "Кадастровый квартал"
            elif re.match(r'^\d{2}:\d{2}:\d{7}:\d{1,5}$', cad_num):
                doc_type = "Выписка ЕГРН"
            else:
                doc_type = "Объект недвижимости"
            
            contract_date = datetime.strptime(row['contract_date'], '%Y-%m-%d')
            date_str = self.format_date_for_doc(contract_date)
            foundation = f"{row['contract_number']} от {date_str}"
            
            row_cells = table.rows[i].cells
            row_cells[0].text = str(i)
            row_cells[1].text = doc_type
            row_cells[2].text = cad_num
            row_cells[3].text = foundation
            
            for cell in row_cells:
                self.set_cell_font_size_12(cell)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Вставляем таблицу после маркера
        marker_paragraph._element.addnext(table._element)
        
        # Удаляем параграф с маркером
        parent.remove(marker_paragraph._element)
    
    def generate_sluzhebka(self, template_path, rows_data, current_date):
        """Генерация служебной записки"""
        # Загружаем шаблон
        doc = Document(template_path)
        
        # Заменяем маркеры дат
        replacements = {
            self.markers['current_date']: self.format_date_for_doc(current_date),
            self.markers['current_date_short']: self.format_date_short(current_date),
        }
        self.replace_in_document(doc, replacements)
        
        # Ищем маркер таблицы
        marker_paragraphs = []
        for paragraph in doc.paragraphs:
            if self.markers['table_place'] in paragraph.text or '##TABLE_ROW##' in paragraph.text:
                marker_paragraphs.append(paragraph)
        
        # Создаем таблицы на месте маркеров
        for marker_paragraph in marker_paragraphs:
            self.create_table_at_marker_sluzhebka(doc, marker_paragraph, rows_data)
        
        # Приводим размер шрифта к 12 для всего текста (кроме заголовка)
        for paragraph in doc.paragraphs:
            self.normalize_paragraph_font_size(paragraph, target_size=12, exclude_title=True)
        
        for table in doc.tables:
            self.normalize_table_font_size(table, target_size=12)
        
        return doc
    
    def generate_kpt(self, template_path, cadastral_number, contract_number, contract_date, current_date):
        """Генерация запроса на кадастровый квартал"""
        doc = Document(template_path)
        
        replacements = {
            self.markers['kadastr_number']: cadastral_number,
            self.markers['contract_number']: contract_number,
            self.markers['contract_date']: self.format_date_for_doc(contract_date),
            self.markers['contract_date_short']: self.format_date_short(contract_date),
            self.markers['current_date']: self.format_date_for_doc(current_date),
            self.markers['current_date_short']: self.format_date_short(current_date),
        }
        
        self.replace_in_document(doc, replacements)
        
        return doc
    
    def generate_zu(self, template_path, cadastral_number, contract_number, contract_date, current_date):
        """Генерация запроса на земельный участок"""
        doc = Document(template_path)
        
        replacements = {
            self.markers['kadastr_number']: cadastral_number,
            self.markers['contract_number']: contract_number,
            self.markers['contract_date']: self.format_date_for_doc(contract_date),
            self.markers['contract_date_short']: self.format_date_short(contract_date),
            self.markers['current_date']: self.format_date_for_doc(current_date),
            self.markers['current_date_short']: self.format_date_short(current_date),
        }
        
        self.replace_in_document(doc, replacements)
        
        return doc
    
    def convert_to_pdf(self, docx_path, pdf_path):
        """Конвертация DOCX в PDF"""
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            return True
        except ImportError:
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)
                doc.Close()
                word.Quit()
                return True
            except:
                return False